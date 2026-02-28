from __future__ import annotations

import re
from typing import List
import pdfplumber
from docx import Document as DocxDocument


def read_pdf_text(path: str) -> str:
    chunks: List[str] = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            txt = p.extract_text() or ""
            chunks.append(txt)
    return "\n".join(chunks)


def read_docx_text(path: str) -> str:
    doc = DocxDocument(path)
    chunks: List[str] = []
    auto_num = 1

    for p in _iter_docx_paragraphs(doc):
        text = (p.text or "").strip()
        if not text:
            continue

        # Keep explicit claim numbering as-is.
        if re.match(r"^\s*\d+[\.\):]\s*", text):
            chunks.append(text)
            m_no = re.match(r"^\s*(\d+)[\.\):]\s*", text)
            if m_no:
                auto_num = max(auto_num, int(m_no.group(1)) + 1)
            continue

        # Word auto-numbered lists often store number metadata, not literal text.
        if _is_numbered_list_paragraph(p):
            chunks.append(f"{auto_num}. {text}")
            auto_num += 1
            continue

        chunks.append(text)

    return "\n".join(chunks)


def _iter_docx_paragraphs(doc: DocxDocument):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _is_numbered_list_paragraph(paragraph) -> bool:
    try:
        ppr = paragraph._p.pPr  # type: ignore[attr-defined]
        if ppr is not None and ppr.numPr is not None:
            return True
    except Exception:
        pass

    style_name = ""
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        style_name = ""

    if "list" in style_name and ("number" in style_name or "num" in style_name):
        return True
    if "numbered" in style_name:
        return True
    return False


def _clean(t: str) -> str:
    t = t.replace("\u00ad", "")
    t = re.sub(r"\(cid:\d+\)", "", t)
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def _extract_amended_claims_from_text(raw_text: str) -> str:
    """
    Extract claim text from amended-claims text.
    - Prefer start at heading variants: WE CLAIM / CLAIMS / REGARDING CLAIMS
    - Fallback to first numbered claim if heading is missing
    - End at FER reply section markers if present
    """
    t = _clean(raw_text or "")
    if not t:
        return ""

    start_idx = 0
    for pat in [
        r"(?im)^\s*WE\s+CLAIM\s*:?\s*$",
        r"(?im)^\s*WE\s+CLAIM\s*:?\s*",
        r"(?im)^\s*CLAIMS?\s*:?\s*$",
        r"(?im)^\s*REGARDING\s+CLAIMS\s*:?\s*$",
        r"\bWe\s+Claim\b\s*:?",
    ]:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            start_idx = m.end()
            break

    tail = t[start_idx:]
    end = len(tail)
    end_markers = [
        r"(?im)^\s*SUBMISSION\s+TO\s+OBJECTION\b",
        r"(?im)^\s*FORMAL\s+REQUIREMENTS\b",
        r"(?im)^\s*YOURS\s+FAITHFULLY\b",
        r"(?im)^\s*ENCLOSURES?\s*:?\s*$",
    ]
    for ep in end_markers:
        m2 = re.search(ep, tail, re.I)
        if m2:
            end = min(end, m2.start())

    claims_block = tail[:end].strip()

    # Claims should normally start with "1." / "1)" / "Claim 1".
    start_pat = r"(?im)^\s*(?:1[\.\):]\s*|Claim\s*1\b)"
    m3 = re.search(start_pat, claims_block)
    if m3:
        claims_block = claims_block[m3.start():].strip()
    else:
        m4 = re.search(start_pat, t)
        if m4:
            tail2 = t[m4.start():]
            end2 = len(tail2)
            for ep in end_markers:
                m5 = re.search(ep, tail2, re.I)
                if m5:
                    end2 = min(end2, m5.start())
            claims_block = tail2[:end2].strip()

    return claims_block


def extract_amended_claims_from_pdf(path: str) -> str:
    return _extract_amended_claims_from_text(read_pdf_text(path))


def extract_amended_claims_from_docx(path: str) -> str:
    return _extract_amended_claims_from_text(read_docx_text(path))
