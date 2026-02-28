from __future__ import annotations

import datetime
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

from .fer_parser import FerParseResult


def _para(
    doc: Document,
    text: str = "",
    bold: bool = False,
    size_pt: int = 11,
    italic: bool = False,
    underline: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        r.font.size = Pt(size_pt)


def _para_red(
    doc: Document,
    text: str = "",
    bold: bool = False,
    size_pt: int = 11,
    italic: bool = False,
    underline: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        r.font.size = Pt(size_pt)
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def _obj_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.size = Pt(11)


def _reply_label(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("OUR REPLY:")
    r.bold = True
    r.font.size = Pt(11)


def _placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _set_cell_placeholder_red(cell, text: str) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p.runs:
        r = p.runs[0]
        r.text = text
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _gap(doc: Document, pts: int = 6) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)


def _iter_cell_paragraphs(cell):
    for p in cell.paragraphs:
        yield p
    for nested in cell.tables:
        for row in nested.rows:
            for c in row.cells:
                yield from _iter_cell_paragraphs(c)


def _justify_document(doc: Document) -> None:
    for p in doc.paragraphs:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in _iter_cell_paragraphs(cell):
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _strip_hindi(text: str) -> str:
    t = re.sub(r"[\u0900-\u097F]+", "", text or "")
    # Preserve original line layout; only normalize excessive spaces per line.
    lines = []
    for ln in t.splitlines():
        ln = re.sub(r"[ \t]{2,}", " ", ln).rstrip()
        lines.append(ln)
    return "\n".join(lines).strip()


def _blocktext(doc: Document, text: str) -> None:
    raw_text = _strip_hindi(text or "")
    if not raw_text:
        return

    merged: List[str] = []

    def flush() -> None:
        if merged:
            _para(doc, " ".join(merged).strip())
            merged.clear()

    for raw_line in raw_text.splitlines():
        s = raw_line.strip()
        if not s:
            flush()
            continue

        latin = re.sub(r"[\u0900-\u097F\s/\-()\[\].,;:0-9]", "", s)
        if len(s) > 3 and len(latin) == 0:
            continue
        if re.match(r"^Page\s+\d+\s+of\s+\d+", s, re.I):
            continue
        if re.match(r"^THE\s+PATENT\s+OFFICE\s*$", s, re.I):
            continue
        if re.match(r"^\([0-9]+\)\.\s*[\u0900-\u097F]", s):
            continue

        if re.match(r"^(?:\(?\d+[.)]|[A-Za-z][.)]|[-*])\s+", s):
            flush()
            _para(doc, s)
            continue

        if s.endswith(":") and len(s) <= 90:
            flush()
            _para(doc, s)
            continue

        merged.append(s)

    flush()


def _extract_numbered_claims(amended_claims: str) -> List[Tuple[int, str]]:
    text = amended_claims or ""
    if not text.strip():
        return []

    pat = re.compile(r"(?im)^\s*(\d+)[\.\):]\s*")
    matches = list(pat.finditer(text))
    if not matches:
        return []

    claims: List[Tuple[int, str]] = []
    for i, m in enumerate(matches):
        no = int(m.group(1))
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block = text[start:end].strip()
        if block:
            claims.append((no, block))
    return claims


def _compact_claim_quote(text: str) -> str:
    t = _strip_hindi(text or "")
    t = t.replace("\r", " ").replace("\n", " ")
    t = re.sub(r"\s+", " ", t).strip()
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)
    t = re.sub(r"\(\s+", "(", t)
    t = re.sub(r"\s+\)", ")", t)
    t = re.sub(r"\[\s+", "[", t)
    t = re.sub(r"\s+\]", "]", t)
    t = re.sub(r'\s+"', '"', t)
    t = re.sub(r'"\s+', '"', t)
    return t


def _normalize_dx_range(dx_range: str) -> str:
    raw = (dx_range or "").strip()
    if not raw:
        return "D1-Dn"

    # Accept common separators and preserve entered range semantics.
    tokens = re.split(r"[,\n;/]+", raw)
    cleaned = []
    for t in tokens:
        s = t.strip().upper()
        if not s:
            continue
        if re.fullmatch(r"D\d+", s):
            cleaned.append(s)
        else:
            cleaned.append(t.strip())

    if not cleaned:
        return "D1-Dn"
    return ", ".join(cleaned)


def _format_d_label_ranges(labels: List[str]) -> str:
    nums = []
    for lbl in labels:
        m = re.fullmatch(r"D(\d{1,3})", (lbl or "").strip().upper())
        if not m:
            return ", ".join([x for x in labels if x]) if labels else ""
        nums.append(int(m.group(1)))

    if not nums:
        return ""

    nums = sorted(set(nums))
    ranges = []
    start = prev = nums[0]
    for n in nums[1:]:
        if n == prev + 1:
            prev = n
            continue
        ranges.append((start, prev))
        start = prev = n
    ranges.append((start, prev))

    parts = []
    for a, b in ranges:
        if a == b:
            parts.append(f"D{a}")
        else:
            parts.append(f"D{a}-D{b}")
    return ", ".join(parts)


def _resolve_dx_display(prior_labels: List[str], dx_range: str) -> str:
    lbls = [x.strip().upper() for x in (prior_labels or []) if (x or "").strip()]
    if lbls:
        ranged = _format_d_label_ranges(lbls)
        if ranged:
            return ranged
        return ", ".join(lbls)
    return _normalize_dx_range(dx_range)


def _normalize_prior_art_entries(prior_art_entries: Optional[List[Dict[str, str]]]) -> List[Dict[str, str]]:
    normalized: List[Dict[str, str]] = []
    for i, row in enumerate(prior_art_entries or [], 1):
        if not isinstance(row, dict):
            continue
        label = (str(row.get("label", "")).strip() or f"D{i}").upper()
        if not re.fullmatch(r"D\d{1,3}", label):
            label = f"D{i}"

        abstract = _strip_hindi(str(row.get("abstract", ""))).strip()
        abstract = re.sub(r"[ \t]{2,}", " ", abstract)
        abstract = re.sub(r"\n{3,}", "\n\n", abstract)

        diagram = _strip_hindi(str(row.get("diagram", ""))).strip()
        diagram = re.sub(r"[ \t]{2,}", " ", diagram)
        diagram_path = str(row.get("diagram_path", "")).strip()

        if not abstract and not diagram and not diagram_path:
            continue
        normalized.append(
            {
                "label": label,
                "abstract": abstract,
                "diagram": diagram,
                "diagram_path": diagram_path,
            }
        )
    return normalized


def _truncate_words(text: str, max_words: int = 80) -> str:
    raw = re.sub(r"\s+", " ", (text or "")).strip()
    if not raw:
        return ""

    words = re.findall(r"\S+", raw)
    if len(words) <= max_words:
        if raw[-1] in ".!?":
            return raw
        end = max(raw.rfind("."), raw.rfind("!"), raw.rfind("?"))
        if end >= int(len(raw) * 0.4):
            return raw[:end + 1].strip()
        return f"{raw}."

    cut = " ".join(words[:max_words]).strip()
    if cut.endswith((".", "!", "?")):
        return cut

    tail_words = words[max_words:max_words + 100]
    if tail_words:
        tail_probe = " ".join(tail_words).strip()
        m = re.search(r"[.!?](?:\s|$)", tail_probe)
        if m:
            return f"{cut} {tail_probe[:m.end()].strip()}".strip()

    back_cut = max(cut.rfind("."), cut.rfind("!"), cut.rfind("?"))
    if back_cut >= int(len(cut) * 0.35):
        return cut[:back_cut + 1].strip()

    return f"{cut}."


def _complete_sentence_text(text: str) -> str:
    raw = re.sub(r"\s+", " ", (text or "")).strip()
    if not raw:
        return ""
    if raw[-1] in ".!?":
        return raw
    return f"{raw}."


def _build_prior_art_disclosure_from_abstracts(prior_arts: List[Dict[str, str]]) -> str:
    lines: List[str] = []
    for row in prior_arts:
        label = row.get("label", "").strip()
        abstract = re.sub(r"\s+", " ", row.get("abstract", "")).strip()
        if not label or not abstract:
            continue
        lines.append(f"{label} discloses {_complete_sentence_text(abstract)}")
    return "\n".join(lines).strip()


def _build_combined_difference_text(
    claim1_text: str,
    prior_arts: List[Dict[str, str]],
    dx_display: str,
) -> str:
    cleaned_claim = re.sub(r"\s+", " ", claim1_text or "").strip()
    if not prior_arts:
        return f"Combined difference over {dx_display}: [INSERT CLAIM-1 VS {dx_display} COMBINED DIFFERENCE ANALYSIS]."

    contrast_parts: List[str] = []
    for row in prior_arts:
        label = (row.get("label", "") or "").strip()
        if not label:
            continue
        abstract = re.sub(r"\s+", " ", row.get("abstract", "")).strip()
        disclosed = _complete_sentence_text(abstract) if abstract else "[INSERT PRIOR-ART ABSTRACT DISCLOSURE]"
        contrast_parts.append(
            f"{label} discloses {disclosed}"
        )

    if not contrast_parts:
        return f"Combined difference over {dx_display}: [INSERT CLAIM-1 VS {dx_display} COMBINED DIFFERENCE ANALYSIS]."

    contrasted = "; ".join(contrast_parts)
    contrast_end = "" if contrasted.endswith((".", "!", "?")) else "."
    return (
        f"Combined difference over {dx_display}: The claimed invention requires the combined feature set of Claim 1 "
        f"({cleaned_claim}). In contrast, {contrasted}{contrast_end} Accordingly, {dx_display} do not individually or in "
        "combination disclose the complete claimed combination."
    )


def _add_prior_art_diagram(doc: Document, diagram_path: str, label: str) -> None:
    path = (diagram_path or "").strip()
    if not path:
        return
    try:
        doc.add_picture(path, width=Inches(5.6))
        _gap(doc, 2)
    except Exception:
        _placeholder(doc, f"[{label} DIAGRAM COULD NOT BE INSERTED]")


def _add_regarding_claims_block(
    doc: Document,
    amended_claims: str,
    dx_range: str = "D1-Dn",
    dx_disclosed_features: str = "",
    prior_art_entries: Optional[List[Dict[str, str]]] = None,
) -> None:
    claims = _extract_numbered_claims(amended_claims)
    if not claims:
        _placeholder(doc, "[INSERT REGARDING-CLAIMS ARGUMENTS HERE]")
        return

    normalized_prior_arts = _normalize_prior_art_entries(prior_art_entries)
    prior_labels = [p["label"] for p in normalized_prior_arts if p.get("label")]
    dx_display = _resolve_dx_display(prior_labels, dx_range)

    dx_features = (dx_disclosed_features or "").strip()
    if not dx_features:
        dx_features = _build_prior_art_disclosure_from_abstracts(normalized_prior_arts)
    if not dx_features:
        dx_features = "[D1-Dn_DISCLOSURE]"

    claim_text_map = {n: txt for n, txt in claims}
    claim1_text = _compact_claim_quote(claim_text_map.get(1, "[INSERT AMENDED CLAIM 1 TEXT]"))
    claim1_line = claim1_text

    _para(doc, "Regarding Claim 1:", bold=True)
    _para(
        doc,
        "The claim 1 is amended to more clearly articulate the subject matter and also to overcome the objections "
        "raised in the first examination report. The amendments are fully supported in the specification on record.",
    )
    _para(
        doc,
        "In determining the differences between the prior art and the claims, the question is not whether the "
        "differences themselves would have been obvious, but whether the claimed invention as a whole would have "
        "been obvious. A prior art reference must be considered in its entirety, as a whole.",
    )
    _para(
        doc,
        "[Emphasis Added] To establish a prima facie case of obviousness, three basic criteria must be met: "
        "(1) there must be some suggestion or motivation to modify the reference or to combine reference teachings; "
        "(2) there must be reasonable expectation of success; and (3) the prior art reference must teach or suggest "
        "all the claim limitations.",
    )
    _para(
        doc,
        f"Thus, Applicant respectfully traverses the rejection because the approach disclosed in {dx_display} and "
        f"the approach claimed in the instant application are not only different, but portions of {dx_display} relied "
        "upon do not render the claimed invention obvious.",
    )
    _para(doc, "Claim 1 has been amended to recite:")
    _para(doc, claim1_text)

    _gap(doc, 2)
    if normalized_prior_arts:
        for row in normalized_prior_arts:
            label = row["label"]
            abstract = row.get("abstract", "")
            if abstract:
                abstract_text = re.sub(r"\s+", " ", abstract).strip()
                if abstract_text:
                    _para(doc, f"{label} discloses {abstract_text}")

            diagram_path = row.get("diagram_path", "").strip()
            if diagram_path:
                _add_prior_art_diagram(doc, diagram_path, label)

            diagram = row.get("diagram", "").strip()
            if diagram and not diagram_path:
                _para(doc, diagram)

        _para_red(doc, _build_combined_difference_text(claim1_text, normalized_prior_arts, dx_display))

        _para(
            doc,
            f"[Emphasis Added] {dx_display} discloses a completely different solution and does not set motivation to combine {dx_display} to arrive at the Applicant claimed invention. Even the problem statement, and the solution of {dx_display} and Applicant claimed invention is different and hence the solutions. The problem statement is clearly evident from background of {dx_display} and Applicant claimed invention. It is to be noted that {dx_display} discloses completely different method and does not disclose the following features of the applicant claimed invention:",
        )
    else:
        _placeholder(doc, f"[INSERT {dx_display} ABSTRACT(S) HERE]")
        _placeholder(doc, f"[EXPLAIN HOW INSTANT INVENTION DIFFERS FROM COMBINED {dx_display}]")

    cmp_table = doc.add_table(rows=2, cols=2)
    cell = cmp_table.rows[1].cells[1]
    cmp_table.style = "Table Grid"
    cmp_table.rows[0].cells[0].text = "Applicant claimed feature"
    cmp_table.rows[0].cells[1].text = f"{dx_display} disclosed features"
    cmp_table.rows[1].cells[0].text = claim1_text
    cmp_table.rows[1].cells[1].text = (
        f"{_strip_hindi(dx_features)}\n\n"
        f"Hence, {dx_display} fail to disclose {claim1_line}."
    )
    cell.add_paragraph("") 
    cell.add_paragraph(
    f"A person with combining skills cannot combine the teachings provided in the prior arts ({dx_display}). "
    f"Hence, {dx_display} fails to disclose the features present in the invention. The interpretation asserted by the examiner is not supported by the cited portions of the {dx_display}. Thus, reconsideration is respectfully requested."
    )
    _para(
            doc,
            f"[Emphasis added] It is important to consider the functions and underlying essence of the invention as described in all steps mentioned in the claims. Therefore, it is respectfully submitted that the interpretation asserted by the Examiner is not supported by the disclosure of {dx_display}. Further, Applicant believe the interpretation asserted by the Examiner regarding the claimed steps is not supported by the disclosure of {dx_display}. Nowhere in the cited portions and the whole document does {dx_display} describe or reasonably suggest the above indicated features claimed in the amended independent claim 1. Therefore, the steps of {dx_display} are different from that of Applicant’s claimed subject matter. Additionally, a prima facie obviousness has not been established. Merely recitation of portions from prior art does not sustain the rejection of obviousness unless the prior art reasonably teaches and provides articulated reasoning with rational underpinning to support the legal conclusion of obviousness. Thus, based on the above, to the extent {dx_display} does not disclose, reasonably teach or suggest the features of the amended independent claim 1, and hence it is respectfully submitted that independent claim 1 is patentable over the cited prior art. Nor does {dx_display} motivate one of ordinary skill in the art to combine {dx_display} with another reference to arrive at the claimed invention. Reconsideration is respectfully requested.",
    )

    for n, txt in claims:
        if n == 1:
            continue
        dep_text = _compact_claim_quote(txt)
        dep_text = re.sub(r"^\s*\d+[\.\):]\s+", "", dep_text).strip()
        dep_text_quoted = dep_text
        if dep_text_quoted:
            if dep_text_quoted.startswith('"') and dep_text_quoted.endswith('"'):
                dep_text_quoted = dep_text_quoted
            else:
                dep_text_quoted = f'"{dep_text_quoted}"'
        _gap(doc, 2)
        _para(doc, f"Regarding Claim {n}:", bold=True)
        _para(
            doc,
            f"Applicant has reviewed the entire application of {dx_display} and found that nowhere in the entire "
            f"applications does {dx_display} describe or reasonably suggest the following features:",
        )
        _para(doc, dep_text_quoted)
        _para(
            doc,
            f"Apart from the above, Applicant believes that dependent claim {n} is allowable not only by virtue of "
            "dependency from patentable independent claim 1, but also by virtue of the additional features the claim "
            "defines.",
        )


_NON_PATENTABILITY_3K_PARA = (
    "Applicant respectfully submits that the subject matter of Claims {n} does not fall within the exclusion of "
    "Section 3(k) of the Patents Act, 1970. The claimed invention represents merely a sequence of algorithmic steps "
    "executed on a conventional computer is inconsistent with the express structural and functional limitations recited "
    "in the claims and fully supported by the Complete Specification. The components in the Applicant claimed invention "
    "are linked to provide a significant technical effect and provides the technical solution to problem. Hence, "
    "applicant preys the Hon. Controller to waive the objection under the section 3(k)."
)

_NON_PATENTABILITY_3M_PARA = (
    "The Applicant further submits that the claims are anchored to specific hardware, do not define gameplay rules or "
    "mental judgments, and are therefore distinguishable from examples falling under Section 3(m). The invention "
    "addresses a technical problem of multi-device data coordination and secure ranking execution, not a method of "
    "playing a game or performing a mental act. All claimed steps are executed by dedicated hardware components, and "
    "the claimed system can be employed in varied ranking or evaluation contexts beyond gaming, underscoring its "
    "industrial applicability. Hence, the applicant prays the Hon. Controller to waive the objection under section 3(m)."
)

_TECH_SOLUTION_LEAD_PARA = (
    "The solution are achieved by providing the technical features that includes technical advancement, contribution "
    "and effect as follows:"
)

_TECH_HARDWARE_FEATURE_PARA = (
    "The intricate hardware features introduced in this invention are expounded upon in the specifications and "
    "corresponding FIGS [INSERT FIG RANGE] to [INSERT FIG RANGE]. Additionally, more comprehensive insights into the "
    "implementation of these unique hardware features can be found in paragraphs [INSERT PARAGRAPH RANGE]."
)

_TECH_3K_REGULATION_PARA = (
    "In accordance with the updated regulations published on June 30, 2017, pertaining to the bestowal of 3k "
    "algorithm/computer-related innovations, the current invention represents a noteworthy technological progression "
    "and is not subject to exclusion under Section 3(k). It comprises ample technical measures and processes that meet "
    "the requirements for being considered a technical advancement."
)

_TECH_CRI_UPDATE_PARA = (
    "Please take note that the Revised Guidelines for Examination of Computer Related Inventions (CRI) that were "
    "issued in 2017 (Page 15) have been updated:"
)

_TECH_CRI_QUOTE_1 = (
    "\"It is well-established that, while establishing patentability, the focus should be on the underlying substance "
    "of the invention and not on the particular form in which it is claimed.\""
)

_TECH_CRI_QUOTE_2 = (
    "\"If in substance, the claim, taken as whole, does not fall in any of the excluded categories, the patent should "
    "not be denied.\""
)

_TECH_PRESENTS_SOLUTION_PARA = (
    "It is worth respectfully noting that the subject matter being claimed presents a solution."
)

_FERID_ALLANI_INTRO_PARA = (
    "In a recent judgment, the Hon'ble Delhi High Court in the case of Ferid Allani Vs Union of India & ORS (Delhi High Court WP(C) 7/2014 & CM APPL 40736/2019), held that:"
)

_FERID_ALLANI_QUOTE_PARA = (
    "\"In today's digital world, when most inventions are based on computer program, it would be retrograde to argue "
    "that all such inventions would not be patentable. Innovations in the field of artificial intelligence, blockchain "
    "technologies and other digital products would be based on computer programs, however the same would not become "
    "non-patentable simply for that reason. Patent applications in these fields would have to be examined to see if "
    "they result in a technical contribution.\""
)

_TECH_EFFECT_BULLET_1 = (
    "- there is bar on patenting is in respect of `computer programs per se` and not all inventions based on computer "
    "programs"
)

_TECH_EFFECT_BULLET_2 = (
    "- claims in a patent application comprising of software/computer programs can have a technical effect and if the "
    "invention (as claimed in the claims) demonstrate a 'technical effect' or a 'technical contribution' (as defined "
    "in the Draft Guidelines for Examination of Computer Related Inventions, 2013, and an excerpt for the same has "
    "been provided below), it is patentable even though it may be based on a computer program."
)

_TECH_EFFECT_GUIDELINE_PARA = (
    "In accordance with the Draft Guidelines for Examination of Computer Related Inventions 2013, technical effect and "
    "technical advancement are defined as follows:"
)

_TECH_EFFECT_DEFINITION_PARA = (
    "For the purposes of these guidelines, a technical effect is defined as a solution to a technical problem that the "
    "invention, as a whole, strives to overcome. Here are a few broad examples of technical effects:"
)

_NON_PATENTABILITY_WRAPUP_PARA = (
    "The Applicant further submits that the proposed claims meet all the necessary requirements under the said Act. "
    "Therefore, the Applicant humbly requests the Learned Controller to kindly consider the proposed claim amendments "
    "and waive the objection raised above."
)


def _claim_numbers_scope_label(claims: List[Tuple[int, str]]) -> str:
    nums = sorted({n for n, _ in claims if n >= 1})
    if not nums:
        return "[INSERT CLAIM NUMBER(S)]"
    if len(nums) == 1:
        return str(nums[0])
    return f"{nums[0]}-{nums[-1]}"


def _extract_claim_text_for_technical_sections(amended_claims: str) -> Tuple[str, List[Tuple[int, str]]]:
    claims = _extract_numbered_claims(amended_claims)
    if not claims:
        raw = _compact_claim_quote(amended_claims)
        raw = re.sub(r"^\s*\d+[\.\):]\s+", "", raw).strip()
        if raw:
            return raw, [(1, raw)]
        return "", []

    by_num = {n: txt for n, txt in claims}
    claim1_raw = by_num.get(1, claims[0][1])
    claim1 = _compact_claim_quote(claim1_raw)
    claim1 = re.sub(r"^\s*\d+[\.\):]\s+", "", claim1).strip()

    claim_entries: List[Tuple[int, str]] = []
    for n, txt in claims:
        body = _compact_claim_quote(txt)
        body = re.sub(r"^\s*\d+[\.\):]\s+", "", body).strip()
        if body:
            claim_entries.append((n, body))
    return claim1, claim_entries


def _claims_single_paragraph(claim_entries: List[Tuple[int, str]]) -> str:
    bodies: List[str] = []
    for _, claim_text in claim_entries:
        body = re.sub(r"\s+", " ", (claim_text or "")).strip()
        body = re.sub(r"^\s*\d+[\.\):]\s+", "", body).strip()
        if body:
            bodies.append(body)
    if not bodies:
        return ""
    merged = " ".join(bodies)
    merged = re.sub(r"\s+([,.;:!?])", r"\1", merged)
    merged = re.sub(r"\s{2,}", " ", merged).strip()
    return merged


def _add_technical_effect_images(doc: Document, image_paths: Optional[List[str]]) -> bool:
    paths = [p for p in (image_paths or []) if (p or "").strip()]
    if not paths:
        return False

    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)

    fig_no = 1
    for path in paths:
        try:
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.add_run().add_picture(path, width=Inches(5.8))

            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.add_run(f"FIG. {fig_no}")
            fig_no += 1

            desc_p = doc.add_paragraph()
            desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = desc_p.add_run("[Enter Description of the diagram]")
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        except Exception:
            _placeholder(doc, "[TECHNICAL EFFECT IMAGE COULD NOT BE INSERTED]")
    return True


def _contains_section_clause(text: str, clause: str) -> bool:
    c = (clause or "").strip().lower()
    if c not in {"k", "m"}:
        return False

    raw = " ".join((text or "").split()).lower()
    if not raw:
        return False

    sec_pat = rf"\bsection\s*3\s*(?:\(\s*{c}\s*\)|{c})(?=\W|$)"
    clause_pat = rf"\bclause\s*\(\s*{c}\s*\)\s*of\s*section\s*\(?\s*3\s*\)?(?=\W|$)"
    short_pat = rf"(?<!\d)3\s*(?:\(\s*{c}\s*\)|{c})(?=\W|$)"
    return bool(re.search(sec_pat, raw) or re.search(clause_pat, raw) or re.search(short_pat, raw))


def _add_non_patentability_static_paras(
    doc: Document,
    objection_text: str,
    claim_scope_label: str,
) -> bool:
    added = False
    if _contains_section_clause(objection_text, "k"):
        _para(doc, _NON_PATENTABILITY_3K_PARA.format(n=claim_scope_label))
        added = True
    if _contains_section_clause(objection_text, "m"):
        _para(doc, _NON_PATENTABILITY_3M_PARA)
        added = True
    return added


def _add_non_patentability_technical_sections(
    doc: Document,
    amended_claims: str,
    cs_background_text: str = "",
    cs_summary_text: str = "",
    cs_technical_effect_text: str = "",
    technical_effect_image_paths: Optional[List[str]] = None,
) -> None:
    bg = (cs_background_text or "").strip()
    sm = (cs_summary_text or "").strip()
    te = (cs_technical_effect_text or "").strip()
    claim1_text, claim_entries = _extract_claim_text_for_technical_sections(amended_claims)

    _gap(doc, 2)
    _obj_label(doc, "TECHNICAL PROBLEM SOLVED BY THE INVENITON:")
    if bg:
        _blocktext(doc, bg)
    else:
        _placeholder(doc, "[INSERT 'BACKGROUND OF THE INVENTION' FROM CS HERE]")

    _gap(doc, 2)
    _obj_label(doc, "TECHNICAL SOLUTION SOLVED BY THE INVENITON:")
    _para(doc, _TECH_SOLUTION_LEAD_PARA, bold=True, underline=True)
    if claim1_text:
        _para(doc, claim1_text)
    else:
        _placeholder(doc, "[INSERT CLAIM-1 FEATURES HERE]")
    _para(doc, _TECH_HARDWARE_FEATURE_PARA, bold=True, underline=True)
    _para(doc, _TECH_3K_REGULATION_PARA)
    _para(doc, _TECH_CRI_UPDATE_PARA)
    _para(doc, _TECH_CRI_QUOTE_1)
    _para(doc, _TECH_CRI_QUOTE_2)
    _para(doc, _TECH_PRESENTS_SOLUTION_PARA)
    if claim1_text:
        _para(doc, claim1_text)
    else:
        _placeholder(doc, "[INSERT CLAIM-1 FEATURES HERE]")
    if sm:
        _blocktext(doc, sm)
    else:
        _placeholder(doc, "[INSERT 'SUMMARY OF THE INVENTION' FROM CS HERE]")

    _para(doc, _FERID_ALLANI_INTRO_PARA)
    _para(doc, _FERID_ALLANI_QUOTE_PARA)
    _para(doc, _TECH_EFFECT_BULLET_1)
    _para(doc, _TECH_EFFECT_BULLET_2)
    _para(doc, _TECH_EFFECT_GUIDELINE_PARA)
    _gap(doc, 2)
    _obj_label(doc, "Technical Effect:")
    _para(doc, _TECH_EFFECT_DEFINITION_PARA)
    te_resolved = te or sm or bg
    if te_resolved:
        _blocktext(doc, te_resolved)
    else:
        _placeholder(doc, "[INSERT TECHNICAL EFFECT HERE]")
    claims_para = _claims_single_paragraph(claim_entries)
    if claims_para:
        _para(doc, claims_para)
    else:
        _placeholder(doc, "[INSERT AMENDED CLAIMS HERE - SINGLE PARAGRAPH, NO NUMBERING]")
    if not _add_technical_effect_images(doc, technical_effect_image_paths):
        _placeholder(doc, "[INSERT TECH_SOLUTION_IMAGES HERE]")
    _para(doc, _NON_PATENTABILITY_WRAPUP_PARA)


_FORMAL_CATEGORY_PATTERNS = [
    ("Form 28", r"Form\s*28\b"),
    ("Form 18", r"Form\s*18\b"),
    ("Form 13", r"Form\s*13\b"),
    ("Form 9", r"Form\s*9\b"),
    ("Form 8", r"Form\s*8\b"),
    ("Form 5", r"Form\s*5\b"),
    ("Form 3", r"Form\s*3\b"),
    ("Form 2", r"Form\s*2\b"),
    ("Form 1", r"Form\s*1\b"),
    ("Stamp Duty", r"Stamp\s+[Dd]uty"),
    ("Power of Attorney", r"Power\s+of\s+Attorney"),
    ("Format of Specification", r"Format\s+of\s+Specification|\(rule\s*13\)"),
    ("Format of Drawings", r"Format\s+of\s+Drawings|In drawings|drawings sheet|section\s*78\(2\)"),
    ("Other Deficiencies", r"Other\s+Deficiencies|fails\s+to\s+comply"),
]


def _category_from_formal_line(line: str) -> Optional[str]:
    s = line.strip(" /:-")
    if not s:
        return None

    # Strong line-level cues first.
    if re.search(r"\bForm\s*28\b", s, re.I):
        return "Form 28"
    if re.search(r"\bForm\s*18\b", s, re.I):
        return "Form 18"
    if re.search(r"\bForm\s*13\b", s, re.I):
        return "Form 13"
    if re.search(r"\bForm\s*9\b", s, re.I):
        return "Form 9"
    if re.search(r"\bForm\s*8\b", s, re.I):
        return "Form 8"
    if re.search(r"\bForm\s*5\b", s, re.I):
        return "Form 5"
    if re.search(r"\bForm\s*3\b", s, re.I):
        return "Form 3"
    if re.search(r"\bForm\s*2\b", s, re.I) and re.search(r"specification|format|provisional|complete", s, re.I):
        return "Form 2"
    if re.search(r"\bForm\s*1\b", s, re.I) and re.search(r"category|serial number|applicant", s, re.I):
        return "Form 1"

    for cat, pat in _FORMAL_CATEGORY_PATTERNS:
        if re.search(r"^(?:In\s+the\s+)?(?:Whether\s+GPA,\s*SPA,)?\s*" + pat, s, re.I):
            return cat
    return None


def _clean_formal_line(line: str) -> str:
    s = (line or "").strip()
    s = re.sub(r"^[/|]+", "", s).strip()
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _clean_formal_remark(remark: str) -> str:
    t = re.sub(r"\s+", " ", (remark or "")).strip()
    if not t:
        return ""

    # Common OCR fragments seen in FER formal tables.
    t = t.replace('words ""', 'words "We Claim"')
    t = re.sub(r"Applicant attention is drawn to of the Patents Act\.?", "Applicant attention is drawn to section 78(2) of the Patents Act.", t, flags=re.I)
    t = re.sub(r"\bto of the Patents Act\.?", "to section 78(2) of the Patents Act.", t, flags=re.I)
    t = re.sub(r"\s+", " ", t).strip()
    t = re.sub(r"\s*-\s*IV\s*:?\s*/?\s*$", "", t, flags=re.I).strip()

    # Deduplicate repeated sentences while preserving order.
    parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", t) if p.strip()]
    dedup = []
    seen = set()
    for p in parts:
        key = re.sub(r"[^a-z0-9]+", "", p.lower())
        if key and key not in seen:
            seen.add(key)
            dedup.append(p)
    if dedup:
        t = " ".join(dedup)
    return t


def _split_mixed_formal_rows(rows: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
    split_cues = [
        ("Form 1", r"\bWhile filing the instant application,\s*in Form\s*1\b"),
        ("Form 2", r"\bIn Form\s*2\b"),
        ("Form 28", r"\bApplicant is required to submit Form 28\b"),
    ]

    out: List[Tuple[str, str]] = []
    for cat, remark in rows:
        r = remark or ""
        split_done = False
        for target_cat, cue in split_cues:
            m = re.search(cue, r, re.I)
            if not m or target_cat == cat:
                continue
            head = _clean_formal_remark(r[:m.start()].strip())
            tail = _clean_formal_remark(r[m.start():].strip())
            if head:
                out.append((cat, head))
            if tail:
                out.append((target_cat, tail))
            split_done = True
            break
        if not split_done:
            out.append((cat, _clean_formal_remark(r)))
    return out


def _parse_formal_rows(text: str):
    if not text:
        return []

    m_hdr = re.search(r"(?:Objections?)[^\n]*?Remarks?", text, re.I)
    table_text = text[m_hdr.end():].strip() if m_hdr else text
    table_text = re.sub(r"\n?Page\s+\d+\s+of\s+\d+\s*\n?THE\s+PATENT\s+OFFICE\s*\n?", "\n", table_text, flags=re.I)
    table_text = re.sub(r"\bPART\s*[-–]\s*IV\b.*$", "", table_text, flags=re.I | re.S).strip()

    lines = [_clean_formal_line(ln) for ln in table_text.splitlines()]
    lines = [ln for ln in lines if ln and len(ln) > 2]

    rows: List[Tuple[str, str]] = []
    current_cat: Optional[str] = None
    current_parts: List[str] = []

    def flush():
        nonlocal current_cat, current_parts
        if not current_cat:
            current_parts = []
            return
        remark = " ".join(current_parts).strip()
        remark = re.sub(r"\s+", " ", remark).strip()
        if remark:
            rows.append((current_cat, _clean_formal_remark(remark)[:1200]))
        current_cat = None
        current_parts = []

    for ln in lines:
        if re.search(r"^DOCUMENTS\s+ON\s+RECORD", ln, re.I):
            break
        if re.search(r"^PART\s*[-–]\s*IV", ln, re.I):
            break

        cat = _category_from_formal_line(ln)
        if cat:
            flush()
            current_cat = cat
            stripped = ln
            for c2, pat in _FORMAL_CATEGORY_PATTERNS:
                if c2 == cat:
                    stripped = re.sub(
                        r"^(?:In\s+the\s+)?(?:Whether\s+GPA,\s*SPA,)?\s*" + pat + r"\s*",
                        "",
                        stripped,
                        flags=re.I,
                    )
                    break
            stripped = stripped.strip(" :-")
            if stripped:
                current_parts.append(stripped)
            continue

        if current_cat:
            current_parts.append(ln)

    flush()

    rows = _split_mixed_formal_rows(rows)

    # Merge duplicate categories while preserving discovery order.
    merged = {}
    order: List[str] = []
    for cat, remark in rows:
        if cat not in merged:
            merged[cat] = []
            order.append(cat)
        merged[cat].append(remark)

    final_rows = []
    for cat in order:
        joined = " ".join(merged[cat])
        joined = re.sub(r"\s+", " ", joined).strip()
        if joined:
            final_rows.append((cat, joined[:900]))

    if not final_rows and table_text:
        return [("Formal Requirements", table_text[:1200])]
    return final_rows


def _add_formal_table(
    doc: Document,
    fer_formal_text: str = "",
    fer_formal_rows: Optional[List[Tuple[str, str]]] = None,
) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    for cell, label in zip(table.rows[0].cells, ["Objections", "Remarks", "Our Reply"]):
        cell.text = label
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    if fer_formal_rows:
        for ob, rem in fer_formal_rows:
            row = table.add_row().cells
            row[0].text = _strip_hindi(ob)
            row[1].text = _strip_hindi(rem)
            _set_cell_placeholder_red(row[2], "[INSERT COMPLIANCE STATEMENT / REPLY HERE]")
    else:
        raw_formal = (fer_formal_text or "").strip()
        if raw_formal:
            row = table.add_row().cells
            row[0].text = "As in FER"
            row[1].text = _strip_hindi(raw_formal)
            _set_cell_placeholder_red(row[2], "[INSERT COMPLIANCE STATEMENT / REPLY HERE]")
        else:
            row = table.add_row().cells
            row[0].text = "[FORMAL OBJECTION CATEGORY]"
            row[1].text = "[PASTE REMARKS FROM FER HERE]"
            _set_cell_placeholder_red(row[2], "[INSERT COMPLIANCE STATEMENT HERE]")


def generate_reply_docx(
    fer: FerParseResult,
    cs_title: str,
    amended_claims: str,
    detailed_obs_text: str = "",
    formal_reqs_text: str = "",
    agent: Optional[str] = None,
    office_address: str = "THE PATENT OFFICE\nI.P.O BUILDING\nG.S.T.Road, Guindy\nChennai - [PIN]",
    dx_range: str = "D1-Dn",
    dx_disclosed_features: str = "",
    prior_art_entries: Optional[List[Dict[str, str]]] = None,
    formal_reqs_rows: Optional[List[Tuple[str, str]]] = None,
    cs_background_text: str = "",
    cs_summary_text: str = "",
    cs_technical_effect_text: str = "",
    technical_effect_image_paths: Optional[List[str]] = None,
) -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    today = datetime.date.today().strftime("%d %B %Y")
    _para(doc, today)
    _para(doc, "To,")

    address_lines = [ln.strip() for ln in (office_address or "").splitlines() if ln.strip()]
    if not address_lines:
        address_lines = ["THE PATENT OFFICE", "I.P.O BUILDING", "G.S.T.Road, Guindy", "Chennai - [PIN]"]
    for ln in address_lines:
        _para(doc, ln)
    _gap(doc, 8)

    controller = fer.controller_name or "[Controller Name]"
    _para(doc, f"Kind Attention: {controller}, Controller of Patents")
    _gap(doc, 4)

    app_no = fer.application_no or "[Application No]"
    filing = fer.filing_date or "[Filing Date]"
    fer_date = fer.fer_dispatch_date or "[FER Date]"
    appl = fer.applicant or "[Applicant]"
    title = cs_title or fer.title or "[Title of Invention]"

    _para(doc, f"Re: Response to FER dated {fer_date}, with respect to Patent Application No: {app_no} filed on {filing}")
    _para(doc, f"Applicant(s): {appl}")
    _para(doc, f'Title: "{title}"')
    _para(doc, f"Letter No: Ref.No/Application No /{app_no} Dated: {fer_date}")
    _gap(doc, 6)
    _para(doc, "Dear Sir,")
    _para(
        doc,
        f"With reference to your letter No Ref/Application No /{app_no} dated {fer_date}, "
        "our humble submissions in the FER matter are as follows for and on behalf of applicant herein:",
    )

    _heading(doc, "AMENDMENTS MADE TO THE CLAIMS ARE AS FOLLOWS")
    _para(doc, "We Claim:", bold=True)
    claims_blocks = _extract_numbered_claims(amended_claims)
    claim_scope_label = _claim_numbers_scope_label(claims_blocks)
    if claims_blocks:
        for _, block in claims_blocks:
            _para(doc, _compact_claim_quote(block))
    else:
        claims_lines = [l.strip() for l in (amended_claims or "").splitlines() if l.strip()]
        if claims_lines:
            _para(doc, _compact_claim_quote("\n".join(claims_lines)))
        else:
            _placeholder(doc, "[PASTE AMENDED CLAIMS HERE - upload the Amended Claims PDF]")
    _gap(doc, 8)

    objections = fer.objections or []
    has_regarding_claims_objection = False
    regarding_claims_content_rendered = False
    non_pat_technical_sections_rendered = False

    if not objections:
        for i in range(1, 7):
            _heading(doc, f"SUBMISSION TO OBJECTION {i}")
            _placeholder(doc, f"[PASTE EXAMINER'S OBJECTION {i} TEXT HERE]")
            _reply_label(doc)
            _placeholder(doc, f"[INSERT REPLY TO OBJECTION {i} HERE]")
    else:
        for obj in objections:
            h = obj.heading.upper()
            if "REGARDING CLAIMS" in h:
                _obj_label(doc, "REGARDING CLAIMS:")
            else:
                _heading(doc, f"SUBMISSION TO OBJECTION {obj.number}")
                _obj_label(doc, obj.heading.upper() + ":")
            _blocktext(doc, obj.body)
            _gap(doc, 4)
            _reply_label(doc)

            if "INVENTIVE STEP" in h:
                if _extract_numbered_claims(amended_claims):
                    _add_regarding_claims_block(
                        doc,
                        amended_claims,
                        dx_range=dx_range,
                        dx_disclosed_features=dx_disclosed_features,
                        prior_art_entries=prior_art_entries,
                    )
                    regarding_claims_content_rendered = True
                    has_regarding_claims_objection = True
                else:
                    _placeholder(doc, "[EXPLAIN HOW AMENDED CLAIM OVERCOMES D1, D2, etc.]")
                    _placeholder(doc, "[ADD INSTANT INVENTION vs PRIOR ART TABLE IF NEEDED]")
            elif "NOVELTY" in h:
                _placeholder(doc, "[INSERT NOVELTY ARGUMENT AGAINST CITED PRIOR ART HERE]")
                _placeholder(doc, "[EXPLAIN DISTINGUISHING FEATURES OF AMENDED CLAIMS HERE]")
            elif "NON PATENTABILITY" in h:
                non_pat_text = f"{obj.heading}\n{obj.body}"
                if not _add_non_patentability_static_paras(doc, non_pat_text, claim_scope_label):
                    _placeholder(doc, "[INSERT SECTION 3(f)/3(o)/3(k) ARGUMENT HERE]")
                _placeholder(doc, "[EXPLAIN WHY INVENTION IS NOT EXCLUDED UNDER CITED CLAUSE]")
                if not non_pat_technical_sections_rendered:
                    _add_non_patentability_technical_sections(
                        doc,
                        amended_claims=amended_claims,
                        cs_background_text=cs_background_text,
                        cs_summary_text=cs_summary_text,
                        cs_technical_effect_text=cs_technical_effect_text,
                        technical_effect_image_paths=technical_effect_image_paths,
                    )
                    non_pat_technical_sections_rendered = True
            elif "REGARDING CLAIMS" in h:
                has_regarding_claims_objection = True
                if regarding_claims_content_rendered:
                    _para(
                        doc,
                        "Detailed claim-wise distinction over cited prior art is already submitted above under the Inventive Step reply.",
                    )
                else:
                    _add_regarding_claims_block(
                        doc,
                        amended_claims,
                        dx_range=dx_range,
                        dx_disclosed_features=dx_disclosed_features,
                        prior_art_entries=prior_art_entries,
                    )
                    regarding_claims_content_rendered = True
            elif "SUFFICIENCY" in h:
                _placeholder(doc, "[INSERT ABSTRACT / SUFFICIENCY COMPLIANCE STATEMENT HERE]")
            elif "CLARITY" in h:
                _placeholder(doc, "[INSERT CLARITY RESPONSE - EXPLAIN HOW AMENDMENTS ADDRESS EACH POINT]")
            elif "DEFINITIVENESS" in h:
                _placeholder(doc, "[INSERT DEFINITIVENESS RESPONSE (Sec 10(4)(c), 10(5)) HERE]")
            elif "SCOPE" in h:
                _placeholder(doc, "[INSERT SCOPE RESPONSE - EXPLAIN HOW CLAIMS DEFINE CLEAR BOUNDARIES]")
            elif "OTHERS" in h:
                _placeholder(doc, "[INSERT RESPONSE TO OTHER REQUIREMENTS HERE]")
            else:
                _placeholder(doc, f"[INSERT REPLY TO OBJECTION {obj.number} HERE]")
            _gap(doc, 8)

    if (
        not has_regarding_claims_objection
        and not regarding_claims_content_rendered
        and _extract_numbered_claims(amended_claims)
    ):
        _obj_label(doc, "REGARDING CLAIMS:")
        _reply_label(doc)
        _add_regarding_claims_block(
            doc,
            amended_claims,
            dx_range=dx_range,
            dx_disclosed_features=dx_disclosed_features,
            prior_art_entries=prior_art_entries,
        )
        regarding_claims_content_rendered = True
        _gap(doc, 8)

    if not non_pat_technical_sections_rendered:
        _heading(doc, "SUBMISSION TO NON PATENTABILITY U/S 3")
        _obj_label(doc, "NON PATENTABILITY U/S 3:")
        _reply_label(doc)
        _placeholder(doc, "[INSERT SECTION 3(k)/3(m) ARGUMENT HERE]")
        _placeholder(doc, "[EXPLAIN WHY INVENTION IS NOT EXCLUDED UNDER CITED CLAUSE]")
        _add_non_patentability_technical_sections(
            doc,
            amended_claims=amended_claims,
            cs_background_text=cs_background_text,
            cs_summary_text=cs_summary_text,
            cs_technical_effect_text=cs_technical_effect_text,
            technical_effect_image_paths=technical_effect_image_paths,
        )
        _gap(doc, 8)

    _heading(doc, "FORMAL REQUIREMENTS:")
    _add_formal_table(doc, formal_reqs_text, formal_reqs_rows)

    _gap(doc, 10)
    _para(
        doc,
        "In the event above submissions are not found to be persuasive, a further hearing/an opportunity for "
        "clarification (through telephone, meeting or the like), preferably in view of Section 80 or Section 14 "
        "may please be granted before taking any adverse decision.",
    )
    _gap(doc, 8)
    _para(doc, "Yours faithfully,")
    _para(doc, "Adv. Pranav Bhat ")
    _para(doc, "(Patent Agent - IN/PA 4580)")
    _gap(doc, 6)
    _para(doc, "Enclosure:")
    _placeholder(doc, "1. [List enclosures here]")

    _justify_document(doc)
    if doc.paragraphs:
        doc.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc
