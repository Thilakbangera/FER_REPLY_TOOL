from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from typing import Dict, List, Tuple

import pdfplumber
from docx import Document as DocxDocument


@dataclass
class PriorArt:
    label: str
    docno: str = ""
    pub_date: str = ""


@dataclass
class Objection:
    number: int
    heading: str
    body: str
    sections: List[str]
    claims: str = ""
    prior_arts: List[PriorArt] = None


@dataclass
class FerParseResult:
    application_no: str = ""
    filing_date: str = ""
    fer_dispatch_date: str = ""
    applicant: str = ""
    title: str = ""
    controller_name: str = ""
    examiner_name: str = ""
    reply_deadline: str = ""
    prior_arts: List[PriorArt] = None
    objections: List[Objection] = None


def read_pdf_text(path: str) -> str:
    chunks: List[str] = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            chunks.append(p.extract_text() or "")
    return "\n".join(chunks)


def read_docx_text(path: str) -> str:
    chunks: List[str] = []
    doc = DocxDocument(path)
    for p in doc.paragraphs:
        chunks.append((p.text or "").strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    chunks.append(txt)
    return "\n".join(chunks)


def _clean(text: str) -> str:
    t = (text or "").replace("\u00ad", "")
    t = re.sub(r"\(cid:\d+\)", "", t)
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def _first_match(pattern: str, text: str, flags: int = 0) -> str:
    m = re.search(pattern, text, flags)
    return m.group(1).strip() if m else ""


def _first_date(text: str) -> str:
    m = re.search(r"([0-9]{2}[-/][0-9]{2}[-/][0-9]{4})", text or "")
    return m.group(1) if m else ""


def _normalize_application_no(raw: str) -> str:
    value = (raw or "").strip().strip("/:-|")
    if not value:
        return ""
    digits = re.sub(r"\D", "", value)
    if len(digits) >= 10:
        return digits
    return re.sub(r"\s+", "", value).upper().strip("/:-|")


def _extract_application_no_from_snippet(snippet: str) -> str:
    s = (snippet or "").replace("|", " ")
    if not s:
        return ""

    patterns = [
        r"\b(\d{10,18})\b",
        r"\b((?:\d[\s-]){10,30}\d)\b",
        r"\b(\d{1,6}\s*/\s*[A-Za-z]{2,10}\s*/\s*\d{2,4})\b",
        r"\b([A-Za-z]{2,10}\s*/\s*[A-Za-z0-9]{3,20}\s*/\s*\d{2,4})\b",
    ]
    for pat in patterns:
        m = re.search(pat, s, re.I)
        if m:
            app_no = _normalize_application_no(m.group(1))
            if app_no:
                return app_no
    return ""


def _looks_like_meta_boundary(line: str) -> bool:
    return bool(re.search(
        r"\b(Application|Date|Dispatch|Filing|Priority|Controller|Examiner|Title|Ref\.?\s*No|"
        r"Letter\s*No|PCT|Patent\s+Office|Report|FER|Claim)\b",
        line,
        re.I,
    ))


def _normalize_applicant_name(text: str) -> str:
    s = (text or "").replace("|", " ").strip(" /:;,.-|")
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"\(cid:\d+\)", "", s)
    s = re.sub(
        r"^(?:Name\s+and\s+Address\s+of\s+the\s+Applicant|Applicants?\s*(?:\(\s*s\s*\))?|Applicant\s*(?:\(\s*s\s*\))?)\s*[:\-]?\s*",
        "",
        s,
        flags=re.I,
    )
    s = re.sub(r"^\(\s*s\s*\)\s*", "", s, flags=re.I)
    s = re.sub(r"^(?:Name\s+Nationality\s+Address)\s*", "", s, flags=re.I)
    s = re.sub(
        r"\s*(?:Request|Exam(?:ination)?|PCT|Date|Filing|Priority|Controller|Examiner)\b.*$",
        "",
        s,
        flags=re.I,
    ).strip()
    s = re.sub(r"\s*(?:Nationality|Address)\s*[:\-].*$", "", s, flags=re.I).strip()
    s = re.sub(r"\s*The\s+following\s+specification\b.*$", "", s, flags=re.I).strip()

    suffix = re.search(
        r"\b(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Company)\b",
        s,
        re.I,
    )
    if suffix:
        s = s[:suffix.end()].strip(" ,;.")

    s = s.strip(" ,;.|")
    if not re.search(r"[A-Za-z]", s):
        return ""
    return s


def _looks_like_address_or_meta(text: str) -> bool:
    s = text or ""
    if not s:
        return False
    if re.search(r"\b(?:Nationality|Address|Name)\b\s*[:\-]?", s, re.I):
        return True
    if re.search(
        r"\b(?:road|street|lane|nagar|city|state|district|post|pin|postcode|building|floor|sector|"
        r"circle|park|block|phase|enclave|miles?|karnataka|telangana|india)\b",
        s,
        re.I,
    ):
        return True
    if re.search(r"\d{3,}", s):
        return True
    return False


def _pick_best_applicant_name(raw_name: str) -> str:
    normalized = _normalize_applicant_name(raw_name)
    if normalized:
        words = re.findall(r"[A-Za-z]+", normalized)
        # Keep clean compact phrases as-is (e.g., "Innovation Centre, Manipal University").
        if words and len(words) <= 14 and not _looks_like_address_or_meta(normalized):
            return normalized

    return (
        _extract_company_name_from_block(raw_name)
        or _extract_institution_name_from_block(raw_name)
        or normalized
        or ""
    )


def _extract_applicant_from_labeled_block(text: str) -> str:
    t = text or ""
    if not t.strip():
        return ""

    stops = (
        r"(?:\n\s*The\s+following\s+specification|\n\s*FIELD\s+OF\s+INVENTION\b|"
        r"\n\s*TECHNICAL\s+FIELD\b|\n\s*BACKGROUND\b|\n\s*OBJECT(?:S|IVE)?\s+OF\s+THE\s+INVENTION\b|\Z)"
    )

    block_match = re.search(
        rf"(?is)\bAPPLICANTS?\s*(?:\(\s*s\s*\))?\s*[:\-]?\s*(.+?){stops}",
        t,
        re.I,
    )
    block = block_match.group(1) if block_match else t

    m_name = re.search(
        rf"(?is)\bName\s*[:\-]\s*(.+?)(?=\n\s*(?:Nationality|Address)\s*[:\-]|{stops})",
        block,
        re.I,
    )
    if not m_name:
        m_name = re.search(
            rf"(?is)\bApplicants?\s*(?:\(\s*s\s*\))?\s*[:\-]?\s*(.+?)(?=\n\s*(?:Nationality|Address)\s*[:\-]?|{stops})",
            block,
            re.I,
        )
    if not m_name:
        return ""

    raw_name = m_name.group(1)
    raw_name = re.sub(r"\n+", " ", raw_name)
    raw_name = re.sub(r"\s+", " ", raw_name).strip(" ,;:-")
    if not raw_name:
        return ""

    candidate = _pick_best_applicant_name(raw_name)
    if not candidate or not re.search(r"[A-Za-z]", candidate):
        return ""
    return candidate


def _extract_applicant_from_cs_tables(path: str) -> str:
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:5]:
                tables = page.extract_tables() or []
                for table in tables:
                    cleaned_rows = []
                    for row in table or []:
                        if not row:
                            continue
                        cleaned_rows.append([_clean_table_cell(c) for c in row])
                    if not cleaned_rows:
                        continue

                    for row in cleaned_rows:
                        row_text = " | ".join([c for c in row if c])
                        m_app = re.search(
                            r"(?is)\bApplicant(?:s|\(\s*s\s*\))?\s*[:\-]?\s*(.+?)(?:\bNationality\b|\bAddress\b|$)",
                            row_text,
                            re.I,
                        )
                        if m_app:
                            raw_name = re.sub(r"\s+", " ", m_app.group(1)).strip(" ,;:-|")
                            if re.fullmatch(r"(?i)(name|nationality|address|indian|applicant(?:s|\(\s*s\s*\))?)", raw_name):
                                raw_name = ""
                            if raw_name:
                                candidate = _pick_best_applicant_name(raw_name)
                                if candidate and re.search(r"[A-Za-z]", candidate):
                                    return candidate

                        m = re.search(r"(?is)\bName\s*[:\-]?\s*(.+?)(?:\bNationality\b|$)", row_text, re.I)
                        if m:
                            raw_name = re.sub(r"\s+", " ", m.group(1)).strip(" ,;:-")
                            if re.fullmatch(r"(?i)(name|nationality|address|indian)", raw_name):
                                continue
                            candidate = _pick_best_applicant_name(raw_name)
                            if candidate and re.search(r"[A-Za-z]", candidate):
                                return candidate

                        for i, cell in enumerate(row):
                            key = (cell or "").strip()
                            if re.fullmatch(r"(?i)applicant(?:s|\(\s*s\s*\))?\s*:?", key):
                                if i + 1 >= len(row):
                                    continue
                                raw_name = re.sub(r"\s+", " ", row[i + 1]).strip(" ,;:-|")
                                if not raw_name:
                                    continue
                                if re.fullmatch(r"(?i)(name|nationality|address|indian|applicant(?:s|\(\s*s\s*\))?)", raw_name):
                                    continue
                                candidate = _pick_best_applicant_name(raw_name)
                                if candidate and re.search(r"[A-Za-z]", candidate):
                                    return candidate

                            if not re.fullmatch(r"(?i)name|name\s*:", (cell or "").strip()):
                                continue
                            if i + 1 >= len(row):
                                continue
                            raw_name = re.sub(r"\s+", " ", row[i + 1]).strip(" ,;:-")
                            if not raw_name:
                                continue
                            if re.fullmatch(r"(?i)(name|nationality|address|indian)", raw_name):
                                continue
                            candidate = _pick_best_applicant_name(raw_name)
                            if candidate and re.search(r"[A-Za-z]", candidate):
                                return candidate
    except Exception:
        return ""

    return ""


def _extract_institution_name_from_block(block: str) -> str:
    s = re.sub(r"\s+", " ", (block or "")).strip()
    if not s:
        return ""

    stop_words = {
        "name", "nationality", "address", "indian",
        "unit", "floor", "sector", "road", "street", "lane", "nagar", "building",
        "tower", "towers", "pin", "postcode", "state", "district", "city",
        "village", "plot", "door", "flat", "apartment", "block", "phase",
        "extension", "enclave", "no", "number", "post", "mile",
        "bengaluru", "bangalore", "hyderabad", "chennai", "karnataka", "telangana", "india",
    }
    suffix_pat = r"(?:University|Institute|College|Academy|School|Laborator(?:y|ies)|Centre|Center|Foundation|Trust|Society|Hospital)"

    candidates = []
    for m in re.finditer(suffix_pat, s, re.I):
        prefix = s[:m.start()]
        prefix_tokens = re.findall(r"[A-Za-z0-9&().'/+-]+", prefix)
        kept = []
        for tok in reversed(prefix_tokens):
            low = tok.lower().strip(".,")
            low_alpha = re.sub(r"[^a-z]", "", low)
            if re.search(r"\d", tok):
                if kept:
                    break
                continue
            if low in stop_words or any(low_alpha.startswith(sw) for sw in stop_words):
                if kept:
                    break
                continue
            kept.append(tok.strip(".,"))
            if len(kept) >= 8:
                break

        if not kept:
            continue

        base = " ".join(reversed(kept)).strip()
        if not base:
            continue

        cand = f"{base} {m.group(0)}"
        cand = _normalize_applicant_name(cand)
        if cand:
            candidates.append(cand)

    if not candidates:
        return ""

    # Prefer the shortest clean institutional name to avoid address spillover.
    candidates = sorted(set(candidates), key=lambda x: (len(x.split()), len(x)))
    return candidates[0]


def _extract_company_name_from_block(block: str) -> str:
    s = re.sub(r"\s+", " ", (block or "")).strip()
    if not s:
        return ""

    stop_words = {
        "unit", "floor", "sector", "road", "street", "lane", "nagar", "building",
        "tower", "towers", "pin", "postcode", "state", "district", "city",
        "village", "plot", "door", "flat", "apartment", "block", "phase",
        "extension", "enclave", "no", "number",
    }
    suffix_pat = r"(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Company)"
    first_token = _first_match(r"^\s*([A-Za-z][A-Za-z0-9&'-]+)", s)

    candidates = []
    for m in re.finditer(suffix_pat, s, re.I):
        prefix = s[:m.start()]
        prefix_tokens = re.findall(r"[A-Za-z0-9&().'/+-]+", prefix)
        kept = []
        for tok in reversed(prefix_tokens):
            low = tok.lower().strip(".,")
            low_alpha = re.sub(r"[^a-z]", "", low)
            if re.search(r"\d", tok):
                if kept:
                    break
                continue
            if low in stop_words or any(low.startswith(sw + "-") for sw in stop_words) or any(low_alpha.startswith(sw) for sw in stop_words):
                if kept:
                    break
                continue
            kept.append(tok.strip(".,"))
            if len(kept) >= 10:
                break

        if not kept:
            continue

        base = " ".join(reversed(kept)).strip()
        if not base:
            continue

        cand = f"{base} {m.group(0)}"
        cand = re.sub(r"\s+", " ", cand).strip(" ,;.")
        cand = re.sub(r"^(?:[A-Z]{1,2}\s+){1,2}", "", cand).strip()

        words = re.findall(r"[A-Za-z]+", cand)
        if len(words) < 2:
            continue
        score = len(words) * 3
        score -= sum(1 for w in words if w.lower() in stop_words) * 4
        if re.search(r"\d", cand):
            score -= 6
        candidates.append((score, cand))

    if not candidates:
        return ""

    best = sorted(candidates, key=lambda x: x[0], reverse=True)[0][1]

    # Reattach leading brand token when OCR dropped it near the suffix.
    if first_token and first_token.lower() not in best.lower():
        if len(best.split()) <= 6:
            best = f"{first_token} {best}"
            best = re.sub(
                rf"^{re.escape(first_token)}\s+[A-Z]{{1,2}}\s+",
                f"{first_token} ",
                best,
            ).strip()

    return _normalize_applicant_name(best)


def _extract_applicant_from_text(text: str) -> str:
    lines = [ln.strip() for ln in text.splitlines()]
    corp_suffix = (
        r"\b(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Company)\b"
    )

    for i, ln in enumerate(lines):
        if not re.search(r"\bApplicant\b", ln, re.I):
            continue

        parts: List[str] = []
        m = re.search(r"\bApplicant\b(?:\s*/\s*[\u0900-\u097F]+)?\s*[:\-]?\s*(.*)$", ln, re.I)
        if m and m.group(1).strip():
            parts.append(m.group(1).strip())

        for nxt in lines[i + 1:i + 6]:
            s = nxt.strip(" /:;-")
            if not s:
                if parts:
                    break
                continue
            if not re.search(r"[A-Za-z]", s):
                continue
            if _looks_like_meta_boundary(s):
                break
            parts.append(s)
            if re.search(corp_suffix, " ".join(parts), re.I):
                break

        candidate = _normalize_applicant_name(" ".join(parts))
        if candidate and len(candidate) > 3:
            return candidate

    for pat in [
        r"(?:^|\n)\s*Applicant\s*[:\-]\s*(.+?)(?:\n|$)",
        r"Applicant\s*/\s*[\u0900-\u097F]+\s*[:\-]?\s*(.+?)(?:\n|$)",
    ]:
        m2 = re.search(pat, text, re.I | re.MULTILINE)
        if m2:
            candidate = _normalize_applicant_name(m2.group(1))
            if candidate and len(candidate) > 3:
                return candidate

    return ""


def _extract_meta(text: str) -> Dict[str, str]:
    t = text

    app_no = ""
    for pat in [
        r"Application\s*No[/.]?\s*[:\-]?\s*([^\n]+)",
        r"Application\s*Number\s*[:\-]?\s*([^\n]+)",
        r"Ref\.?\s*No[^\n]*?Application\s*No[/.]?\s*/?\s*([^\n]+)",
        r"[\u0900-\u097F]+[^\n]*?Application\s*No[/.]?\s*/?\s*([^\n]+)",
    ]:
        m = re.search(pat, t, re.I)
        if m:
            app_no = _extract_application_no_from_snippet(m.group(1))
            if app_no:
                break

    if not app_no:
        for ln in t.splitlines()[:40]:
            if re.search(r"Application\s*(?:No|Number)|Ref\.?\s*No", ln, re.I):
                app_no = _extract_application_no_from_snippet(ln)
                if app_no:
                    break

    if not app_no:
        m = re.search(r"\b(\d{12})\b", "\n".join(t.splitlines()[:30]))
        if m:
            app_no = m.group(1)

    if not app_no:
        m = re.search(r"\b(\d{10,18})\b", t)
        if m:
            app_no = m.group(1)

    app_no = _normalize_application_no(app_no)

    filing = ""
    lines = t.splitlines()
    filing_keys = [r"\bDate\s*of\s*Filing\b", r"\bFiling\s*Date\b"]

    # Prefer explicit "Date of Filing" lines in FER metadata region.
    for i, ln in enumerate(lines[:120]):
        if any(re.search(k, ln, re.I) for k in filing_keys):
            filing = _first_date(ln)
            if not filing:
                filing = _first_date(" ".join(lines[i:i + 3]))
            if filing:
                break

    if not filing:
        filing = _first_match(
            r"(?:Date\s*of\s*Filing|Filing\s*Date)\s*[:\-]?\s*([0-9]{2}[-/][0-9]{2}[-/][0-9]{4})",
            t,
            re.I,
        )

    applicant = _extract_applicant_from_text(t)

    fer_date = _first_match(
        r"Date\s*of\s*Dispatch(?:/Email)?\s*[:\-]?\s*([0-9]{2}[-/][0-9]{2}[-/][0-9]{4})",
        t,
        re.I,
    )

    controller = _first_match(r"\n\s*([A-Z][A-Za-z .]+)\s*\n\s*Controller\s+of\s+Patents\b", "\n" + t, re.I)
    if not controller:
        controller = _first_match(
            r"(?:Name\s*of\s*the\s*Controller|Controller.*?Name)\s*[:\-]?\s*([A-Z][A-Za-z .]+)",
            t,
            re.I,
        )

    examiner = _first_match(r"Name\s*of\s*the\s*Examiner\s*[:\-]?\s*([A-Z][A-Za-z .]+)", t, re.I)

    deadline = _first_match(
        r"Last\s*date\s*for\s*filing\s*response(?:\s*to\s*the\s*Examination\s*Report)?\s*[:\-]?\s*([0-9]{2}[-/][0-9]{2}[-/][0-9]{4})",
        t,
        re.I,
    )

    return {
        "application_no": app_no,
        "filing_date": filing,
        "fer_dispatch_date": fer_date,
        "applicant": applicant,
        "title": "",
        "controller_name": controller,
        "examiner_name": examiner,
        "reply_deadline": deadline,
    }


def extract_detailed_observations_block(text: str) -> str:
    t = _clean(text)
    for sp in [
        r"B\.\s*Detailed\s+observations\s+on\s+the\s+requirements\s+under\s+the\s+Act",
        r"Detailed\s+observations\s+on\s+the\s+requirements\s+under\s+the\s+Act",
    ]:
        m = re.search(sp, t, re.I)
        if not m:
            continue
        start_idx = m.start()
        tail = t[start_idx:]
        end_idx = len(tail)
        for ep in [
            r"PART\s*[-–]\s*III\s*[:\-]\s*FORMAL",
            r"PART\s*[-–]\s*III",
            r"FORMAL\s+REQUIREMENTS",
        ]:
            m2 = re.search(ep, tail, re.I)
            if m2:
                end_idx = min(end_idx, m2.start())
        return tail[:end_idx].strip()
    return ""


def extract_formal_requirements_block(text: str) -> str:
    """Extract raw PART-III formal-requirements text from FER."""
    # Keep raw line layout for table fidelity; only remove hard extraction noise.
    t = (text or "").replace("\u00ad", "")
    t = re.sub(r"\(cid:\d+\)", "", t)
    starts = []

    start_patterns = [
        r"PART\s*[-–]\s*III\s*[:\-]\s*FORMAL\s+REQUIREMENTS",
        r"PART\s*[-–]\s*III[^\n]{0,100}FORMAL\s+REQUIREMENTS",
        r"(?m)^\s*FORMAL\s+REQUIREMENTS\s*$",
    ]
    for pat in start_patterns:
        starts.extend(list(re.finditer(pat, t, re.I | re.MULTILINE)))

    if not starts:
        return ""

    # Use the latest occurrence to avoid intro text like "... formal requirements and documents on record."
    start_match = max(starts, key=lambda m: m.start())
    tail = t[start_match.end():]

    end_idx = len(tail)
    for ep in [r"PART\s*[-–]\s*IV", r"DOCUMENTS\s+ON\s+RECORD"]:
        m2 = re.search(ep, tail, re.I)
        if m2:
            end_idx = min(end_idx, m2.start())

    return tail[:end_idx].strip()


def _clean_table_cell(cell: str) -> str:
    t = (cell or "").replace("\u00ad", "")
    t = re.sub(r"\(cid:\d+\)", "", t)
    lines = []
    for ln in t.splitlines():
        ln = re.sub(r"[ \t]+", " ", ln).strip()
        if ln:
            lines.append(ln)
    return "\n".join(lines).strip()


def extract_formal_requirements_rows_from_pdf(path: str) -> List[Tuple[str, str]]:
    """
    Extract formal requirements as row pairs: (Objection, Remark)
    from PART-III table in FER PDF.
    """
    rows: List[Tuple[str, str]] = []
    stop_scan = False

    with pdfplumber.open(path) as pdf:
        in_formal = False
        seen_header = False

        for page in pdf.pages:
            if stop_scan:
                break
            page_text = page.extract_text() or ""

            if not in_formal:
                if re.search(r"PART\s*[-–]?\s*III[^\n]{0,100}FORMAL\s+REQUIREMENTS", page_text, re.I):
                    in_formal = True
                else:
                    continue

            tables = page.extract_tables() or []
            for table in tables:
                if stop_scan:
                    break
                cleaned_rows = []
                for row in table or []:
                    if not row:
                        continue
                    cleaned_rows.append([_clean_table_cell(c) for c in row])
                if not cleaned_rows:
                    continue

                # Skip obvious non-formal tables such as docket/document lists.
                head_text = " ".join(" ".join(r).lower() for r in cleaned_rows[:3])
                if re.search(r"docket|entry number|publication date|sl\.?no", head_text, re.I):
                    continue

                ob_idx = rem_idx = None
                start_idx = 0
                # Some scans include a long merged paragraph containing both words
                # "Objections" and "Remarks" in one cell before the actual header row.
                # Treat a row as header only when both labels appear in different cells.
                for ridx, row in enumerate(cleaned_rows[:8]):
                    ob_cols = [idx for idx, val in enumerate(row) if re.search(r"\bobjections?\b", val, re.I)]
                    rem_cols = [idx for idx, val in enumerate(row) if re.search(r"\bremarks?\b", val, re.I)]
                    chosen = None
                    for oi in ob_cols:
                        for ri in rem_cols:
                            if oi != ri:
                                chosen = (oi, ri)
                                break
                        if chosen:
                            break
                    if chosen:
                        ob_idx, rem_idx = chosen
                        start_idx = ridx + 1
                        seen_header = True
                        break

                # If header is not found but we've already entered formal section,
                # treat this as continuation table and infer likely text columns.
                if ob_idx is None or rem_idx is None:
                    if not seen_header:
                        continue
                    objection_hint = re.search(
                        r"\b(Form\s*\d+|Power\s+of\s+Attorney|Format\s+of|Other\s+Deficiencies|"
                        r"Applicable\s+fee|Endorsement|Date\s+and\s+Signature|Statement\s*&\s*Under\s*Taking|"
                        r"Statement\s*&\s*Undertaking)\b",
                        " ".join(" ".join(r) for r in cleaned_rows),
                        re.I,
                    )
                    if not objection_hint:
                        continue

                    ncols = max((len(r) for r in cleaned_rows), default=0)
                    if ncols < 2:
                        continue
                    col_scores = []
                    for ci in range(ncols):
                        col_vals = [r[ci] if ci < len(r) else "" for r in cleaned_rows]
                        non_empty = [v for v in col_vals if v]
                        score = len(non_empty)
                        avg_len = (sum(len(v) for v in non_empty) / score) if score else 0
                        col_scores.append((score, avg_len, ci))
                    if len(col_scores) < 2:
                        continue

                    # Objection column tends to be shorter text than remarks.
                    ob_candidates = [x for x in col_scores if x[0] > 0]
                    ob_candidates.sort(key=lambda x: (x[1], -x[0]))
                    ob_idx = ob_candidates[0][2]

                    rem_candidates = [x for x in col_scores if x[0] > 0 and x[2] != ob_idx]
                    # In most FER continuation tables, remarks sit to the right of objections.
                    # Prefer right-side columns to avoid noisy merged left columns from OCR.
                    right_candidates = [x for x in rem_candidates if x[2] > ob_idx]
                    if right_candidates:
                        right_candidates.sort(key=lambda x: (-x[0], -x[1], x[2]))
                        rem_idx = right_candidates[0][2]
                    else:
                        rem_candidates.sort(key=lambda x: (-x[1], -x[0]))
                        rem_idx = rem_candidates[0][2] if rem_candidates else None
                    if rem_idx is None:
                        continue
                    start_idx = 0

                for row in cleaned_rows[start_idx:]:
                    ob = row[ob_idx] if ob_idx < len(row) else ""
                    rem = row[rem_idx] if rem_idx < len(row) else ""

                    if not rem:
                        # Some scans place remarks in another non-objection cell.
                        for ci, val in enumerate(row):
                            if ci == ob_idx:
                                continue
                            if val and len(val) > len(rem):
                                rem = val

                    ob = ob.strip()
                    rem = rem.strip()
                    combined = f"{ob} {rem}".strip()

                    if not ob and not rem:
                        continue
                    if re.search(r"\bobjections?\b", combined, re.I) and re.search(r"\bremarks?\b", combined, re.I):
                        continue
                    if re.fullmatch(r"Page", rem, re.I):
                        continue
                    if re.search(r"^THE\s+PATENT\s+OFFICE$", rem, re.I) and not ob:
                        continue
                    if re.search(r"^THE\s+PATENT\s+OFFICE$", ob, re.I) and re.search(r"^THE\s+PATENT\s+OFFICE$", rem, re.I):
                        continue
                    if re.search(r"^Page\s+\d+\s+of\s+\d+$", ob, re.I) or re.search(r"^Page\s+\d+\s+of\s+\d+$", rem, re.I):
                        continue
                    if re.search(r"^THE\s+PATENT\s+OFFICE$", ob, re.I) and not rem:
                        continue

                    combined = f"{ob} {rem}".strip()
                    if re.search(r"PART\s*[-–]?\s*IV|DOCUMENTS\s+ON\s+RECORD", combined, re.I):
                        part_split = re.split(r"PART\s*[-–]?\s*IV|DOCUMENTS\s+ON\s+RECORD", rem, maxsplit=1, flags=re.I)
                        rem_before = part_split[0].strip() if part_split else ""
                        if rem_before and rows:
                            prev_ob, prev_rem = rows[-1]
                            sep = "\n" if prev_rem else ""
                            rows[-1] = (prev_ob, f"{prev_rem}{sep}{rem_before}".strip())
                        stop_scan = True
                        break

                    if ob and rem:
                        rows.append((ob, rem))
                    elif ob and not rem:
                        rows.append((ob, ""))
                    elif rem and rows:
                        prev_ob, prev_rem = rows[-1]
                        sep = "\n" if prev_rem else ""
                        rows[-1] = (prev_ob, f"{prev_rem}{sep}{rem}".strip())

            if stop_scan:
                break
            if in_formal and re.search(r"PART\s*[-–]?\s*IV|DOCUMENTS\s+ON\s+RECORD", page_text, re.I):
                break

    cleaned_pairs: List[Tuple[str, str]] = []
    for ob, rem in rows:
        ob2 = re.sub(r"\bPage\s+\d+\s+of\s+\d+\b", "", ob, flags=re.I)
        rem2 = re.sub(r"\bPage\s+\d+\s+of\s+\d+\b", "", rem, flags=re.I)
        ob2 = re.sub(r"[\u0900-\u097F]+", "", ob2)
        rem2 = re.sub(r"[\u0900-\u097F]+", "", rem2)
        ob2 = re.sub(r"\bTHE\s+PATENT\s+OFFICE\b", "", ob2, flags=re.I).strip(" -:/\n\t")
        rem2 = re.sub(r"\bTHE\s+PATENT\s+OFFICE\b", "", rem2, flags=re.I).strip(" -:/\n\t")
        ob2 = re.sub(r"^\s*/?\s*Objections?\s*/?\s*Remarks?\s*$", "", ob2, flags=re.I)
        rem2 = re.sub(r"^\s*/?\s*Objections?\s*/?\s*Remarks?\s*$", "", rem2, flags=re.I)
        rem2 = re.sub(r"(?im)^\s*[-/]*\s*(?:PART\s*[-–—]?\s*)?IV\s*:?\s*$", "", rem2)
        rem2 = re.sub(r"\s*[-–—]?\s*IV\s*:?\s*$", "", rem2, flags=re.I).strip(" -:/\n\t")
        ob2 = re.sub(r"[ \t]{2,}", " ", ob2).strip()
        rem2 = "\n".join(ln.strip() for ln in rem2.splitlines() if ln.strip())
        if not ob2 and not rem2:
            continue
        cleaned_pairs.append((ob2, rem2))

    # Remove exact duplicate contiguous rows.
    deduped: List[Tuple[str, str]] = []
    for r in cleaned_pairs:
        if deduped and deduped[-1] == r:
            continue
        deduped.append(r)

    # Merge contiguous entries with the same objection label.
    merged: List[Tuple[str, str]] = []
    for ob, rem in deduped:
        if merged and merged[-1][0].strip().lower() == ob.strip().lower():
            prev_ob, prev_rem = merged[-1]
            sep = "\n" if prev_rem and rem else ""
            merged[-1] = (prev_ob, f"{prev_rem}{sep}{rem}".strip())
        else:
            merged.append((ob, rem))

    return merged


def _extract_prior_arts(text: str) -> List[PriorArt]:
    arts: Dict[str, PriorArt] = {}
    for m in re.finditer(
        r"\b(D\d{1,3})\s*[:\-]\s*([A-Z]{2}[A-Z0-9]{4,})\s*(?:\(|Pub\s*Date\s*[:\-]?\s*)?([0-9]{2}[-/][0-9]{2}[-/][0-9]{4})",
        text,
        re.I,
    ):
        lab = m.group(1).upper()
        if lab not in arts:
            arts[lab] = PriorArt(label=lab, docno=m.group(2), pub_date=m.group(3))
    return sorted(arts.values(), key=lambda a: a.label)


def _normalize_heading(raw: str) -> str:
    h = re.sub(r"\s+", " ", (raw or "").upper()).strip()
    h = h.replace("NON-PATENTABILITY", "NON PATENTABILITY")
    if re.fullmatch(r"SCOPE(?:\s+OF(?:\s+THE)?\s+CLAIMS?)?", h, re.I):
        return "SCOPE"
    if h.startswith("OTHER REQUIREMENT") or h.startswith("OTHERS REQUIREMENT"):
        return "OTHERS REQUIREMENTS"
    return h


def _split_objections(text: str) -> List[Tuple[str, str]]:
    t = _clean(text)
    if not t:
        return []

    m_formal = re.search(r"(?:PART\s*[-–]\s*III|FORMAL\s+REQUIREMENTS)", t, re.I)
    if m_formal:
        t = t[:m_formal.start()]

    heading_pat = (
        r"(?P<head>NOVELTY|INVENTIVE STEP|NON[\s\-]PATENTABILITY|REGARDING CLAIMS|"
        r"SUFFICIENCY OF DISCLOSURE|CLARITY AND CONCISENESS|DEFINITIVENESS|"
        r"SCOPE(?:\s+OF(?:\s+THE)?\s+CLAIMS?)?|OTHERS?\s+REQUIREMENTS?)"
    )
    splitter_strict = rf"(?im)^\s*(?:\(\d+\)\.)?\s*/?\s*{heading_pat}\s*[:\-]?\s*$"
    splitter_fallback = rf"(?i){heading_pat}\s*:"

    matches = list(re.finditer(splitter_strict, t))
    if not matches:
        matches = list(re.finditer(splitter_fallback, t))
    if not matches:
        return []

    parts: List[Tuple[str, str]] = []
    for i, m in enumerate(matches):
        head = _normalize_heading(m.group("head"))
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(t)
        body = t[start:end].strip()
        if body:
            parts.append((head, body))
    return parts


def _sections_from_text(body: str) -> List[str]:
    secs = set()
    for m in re.finditer(r"\b(\d+\(\d+\)\([a-z]\)|\d+\([a-z]\)|\d+\(\d+\))\b", body, re.I):
        secs.add(m.group(1))
    for m in re.finditer(r"\b(3\([a-z]\))\b", body, re.I):
        secs.add(m.group(1))
    for m in re.finditer(r"\bRule\s*\d+\s*\(?\d*\)?", body, re.I):
        secs.add(m.group(0).strip())
    return sorted(secs)


def extract_title_from_cs_pdf(path: str) -> str:
    raw = read_pdf_text(path)
    if not raw:
        return ""

    lines = [ln.strip() for ln in raw.splitlines()]

    def _clean_title_line(s: str) -> str:
        x = (s or "").strip()
        x = re.sub(r"\(cid:\d+\)", "", x)
        x = re.sub(r"^\[\d{1,4}\]\s*", "", x)
        x = re.sub(r"^\d+\s+", "", x)
        x = re.sub(r"\s+", " ", x).strip(" :-")
        return x

    stop_pat = re.compile(
        r"\b(NAME\s+AND\s+ADDRESS\s+OF\s+THE\s+APPLICANT|APPLICANTS?|APPLICANT|NATIONALITY|ADDRESS|"
        r"TECHNICAL\s+FIELD|FIELD\s+OF\s+INVENTION|BACKGROUND|OBJECT\s+OF\s+THE\s+INVENTION|"
        r"SUMMARY\s+OF\s+THE\s+INVENTION|DETAILED\s+DESCRIPTION|CLAIMS?|ABSTRACT)\b",
        re.I,
    )

    # Prefer explicit key-value title rows from CS tables when present.
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:5]:
                tables = page.extract_tables() or []
                for table in tables:
                    cleaned_rows = []
                    for row in table or []:
                        if not row:
                            continue
                        cleaned_rows.append([_clean_table_cell(c) for c in row])
                    for row in cleaned_rows:
                        if not row:
                            continue
                        for ci, cell in enumerate(row):
                            c = (cell or "").strip()
                            if not c:
                                continue
                            m_inline = re.search(r"(?is)\bTitle\s*[:\-]\s*(.+)$", c, re.I)
                            if m_inline:
                                cand = _clean_title_line(m_inline.group(1))
                                if cand and not stop_pat.search(cand):
                                    return cand

                            if not re.fullmatch(r"(?i)title\s*:?", c):
                                continue
                            others = [x for j, x in enumerate(row) if j != ci and (x or "").strip()]
                            if not others:
                                continue
                            cand = _clean_title_line(" ".join(others))
                            if cand and not stop_pat.search(cand):
                                return cand
    except Exception:
        pass

    for i, ln in enumerate(lines):
        if not re.search(r"\bTITLE\s+OF\s+THE\s+INVENTION\b", ln, re.I):
            continue

        inline = _clean_title_line(re.sub(r"^.*?\bTITLE\s+OF\s+THE\s+INVENTION\b\s*[:\-]?\s*", "", ln, flags=re.I))
        if inline and not stop_pat.search(inline):
            return inline

        parts: List[str] = []
        for nxt in lines[i + 1:i + 8]:
            s = _clean_title_line(nxt)
            if not s:
                if parts:
                    break
                continue
            if re.match(r"^Page\s+\d+\s+of\s+\d+$", s, re.I):
                continue
            if stop_pat.search(s):
                break
            if re.fullmatch(r"FORM\s*\d+.*", s, re.I):
                continue
            parts.append(s)
            if len(" ".join(parts)) >= 140:
                break

        title = re.sub(r"\s+", " ", " ".join(parts)).strip(" :-")
        if title:
            return title

    t = _clean(raw)
    m2 = re.search(r"(?:^|\n)\s*Title\s*[:\-]\s*([A-Za-z][A-Za-z0-9 &/\',.-]{5,})", t, re.I)
    if m2:
        return m2.group(1).strip()
    return ""


def extract_title_from_cs_docx(path: str) -> str:
    raw = read_docx_text(path)
    if not raw:
        return ""

    lines = [ln.strip() for ln in raw.splitlines()]

    def _clean_title_line(s: str) -> str:
        x = (s or "").strip()
        x = re.sub(r"\(cid:\d+\)", "", x)
        x = re.sub(r"^\[\d{1,4}\]\s*", "", x)
        x = re.sub(r"^\d+\s+", "", x)
        x = re.sub(r"\s+", " ", x).strip(" :-")
        return x

    stop_pat = re.compile(
        r"\b(NAME\s+AND\s+ADDRESS\s+OF\s+THE\s+APPLICANT|APPLICANTS?|APPLICANT|NATIONALITY|ADDRESS|"
        r"TECHNICAL\s+FIELD|FIELD\s+OF\s+INVENTION|BACKGROUND|OBJECT\s+OF\s+THE\s+INVENTION|"
        r"SUMMARY\s+OF\s+THE\s+INVENTION|DETAILED\s+DESCRIPTION|CLAIMS?|ABSTRACT)\b",
        re.I,
    )

    for i, ln in enumerate(lines):
        if not re.search(r"\bTITLE\s+OF\s+THE\s+INVENTION\b", ln, re.I):
            continue

        inline = _clean_title_line(re.sub(r"^.*?\bTITLE\s+OF\s+THE\s+INVENTION\b\s*[:\-]?\s*", "", ln, flags=re.I))
        if inline and not stop_pat.search(inline):
            return inline

        parts: List[str] = []
        for nxt in lines[i + 1:i + 8]:
            s = _clean_title_line(nxt)
            if not s:
                if parts:
                    break
                continue
            if re.match(r"^Page\s+\d+\s+of\s+\d+$", s, re.I):
                continue
            if stop_pat.search(s):
                break
            if re.fullmatch(r"FORM\s*\d+.*", s, re.I):
                continue
            parts.append(s)
            if len(" ".join(parts)) >= 140:
                break

        title = re.sub(r"\s+", " ", " ".join(parts)).strip(" :-")
        if title:
            return title

    t = _clean(raw)
    m2 = re.search(r"(?:^|\n)\s*Title\s*[:\-]\s*([A-Za-z][A-Za-z0-9 &/\',.-]{5,})", t, re.I)
    if m2:
        return m2.group(1).strip()
    return ""


def extract_applicant_from_cs_pdf(path: str) -> str:
    raw = read_pdf_text(path)
    if not raw:
        return ""

    by_label = _extract_applicant_from_labeled_block(raw)
    if by_label:
        return by_label

    t = _clean(raw)
    lines = [ln.strip() for ln in raw.splitlines()]
    start = None

    for i, ln in enumerate(lines):
        if re.search(r"NAME\s+AND\s+ADDRESS\s+OF\s+THE\s+APPLICANT", ln, re.I):
            start = i
            break

    if start is not None:
        parts: List[str] = []
        for ln in lines[start + 1:start + 20]:
            s = re.sub(r"\(cid:\d+\)", "", ln or "").strip(" /:;-")
            if not s:
                if parts:
                    break
                continue
            if re.match(r"^Page\s+\d+\s+of\s+\d+$", s, re.I):
                continue
            if re.search(r"\bName\s+Nationality\s+Address\b", s, re.I):
                continue
            if re.search(r"\bThe\s+following\s+specification\b", s, re.I):
                break
            if not re.search(r"[A-Za-z]", s):
                continue
            if re.search(
                r"\b(NAME\s+AND\s+ADDRESS|NATIONALITY|TITLE\s+OF\s+THE\s+INVENTION|FIELD\s+OF\s+INVENTION|BACKGROUND)\b",
                s,
                re.I,
            ):
                break
            parts.append(s)
            if re.search(
                r"\b(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Company)\b",
                " ".join(parts),
                re.I,
            ):
                break

        raw_block = " ".join(parts)
        candidate = _pick_best_applicant_name(raw_block)
        if candidate:
            return candidate

    by_table = _extract_applicant_from_cs_tables(path)
    if by_table:
        return by_table

    fallback = _extract_applicant_from_text(t)
    candidate = _pick_best_applicant_name(fallback)
    return candidate


def extract_applicant_from_cs_docx(path: str) -> str:
    raw = read_docx_text(path)
    if not raw:
        return ""

    by_label = _extract_applicant_from_labeled_block(raw)
    if by_label:
        return by_label

    t = _clean(raw)
    lines = [ln.strip() for ln in raw.splitlines()]
    start = None

    for i, ln in enumerate(lines):
        if re.search(r"NAME\s+AND\s+ADDRESS\s+OF\s+THE\s+APPLICANT", ln, re.I):
            start = i
            break

    if start is not None:
        parts: List[str] = []
        for ln in lines[start + 1:start + 20]:
            s = re.sub(r"\(cid:\d+\)", "", ln or "").strip(" /:;-")
            if not s:
                if parts:
                    break
                continue
            if re.match(r"^Page\s+\d+\s+of\s+\d+$", s, re.I):
                continue
            if re.search(r"\bName\s+Nationality\s+Address\b", s, re.I):
                continue
            if re.search(r"\bThe\s+following\s+specification\b", s, re.I):
                break
            if not re.search(r"[A-Za-z]", s):
                continue
            if re.search(
                r"\b(NAME\s+AND\s+ADDRESS|NATIONALITY|TITLE\s+OF\s+THE\s+INVENTION|FIELD\s+OF\s+INVENTION|BACKGROUND)\b",
                s,
                re.I,
            ):
                break
            parts.append(s)
            if re.search(
                r"\b(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s*Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?|Company)\b",
                " ".join(parts),
                re.I,
            ):
                break

        raw_block = " ".join(parts)
        candidate = _pick_best_applicant_name(raw_block)
        if candidate:
            return candidate

    fallback = _extract_applicant_from_text(t)
    candidate = _pick_best_applicant_name(fallback)
    return candidate


def _find_first_match(patterns: List[str], text: str, start: int = 0):
    best = None
    for pat in patterns:
        m = re.search(pat, text[start:], re.I | re.M)
        if not m:
            continue
        abs_start = start + m.start()
        if best is None or abs_start < best[0]:
            best = (abs_start, start + m.end(), m)
    return best


def _clean_cs_section_text(section: str) -> str:
    text = (section or "").replace("\u00ad", "")
    text = re.sub(r"\(cid:\d+\)", "", text)
    text = text.replace("\r", "\n")

    cleaned_lines: List[str] = []
    for raw in text.splitlines():
        line = re.sub(r"[ \t]+", " ", raw or "").strip()
        if not line:
            if cleaned_lines and cleaned_lines[-1]:
                cleaned_lines.append("")
            continue

        if re.match(r"^Page\s+\d+\s+of\s+\d+$", line, re.I):
            continue
        if re.match(r"^\d+\s*\|\s*Page\b", line, re.I):
            continue
        if re.match(r"^Page\s+\d+$", line, re.I):
            continue
        if re.match(r"^THE\s+PATENT\s+OFFICE$", line, re.I):
            continue
        if re.match(r"^[\[\(]?\d{1,4}[\]\)]?$", line):
            continue

        # Remove paragraph/line numbering prefixes from CS OCR.
        line = re.sub(r"^\[\d{3,5}\]\s*", "", line)
        line = re.sub(r"^\(?\d{1,3}\)?\s+(?=[A-Za-z])", "", line)
        line = re.sub(r"^\(?\d{1,3}\)?[.:]\s*", "", line)
        line = re.sub(r"\s{2,}", " ", line).strip()

        if not line:
            continue
        cleaned_lines.append(line)

    out_lines: List[str] = []
    for ln in cleaned_lines:
        if ln == "" and out_lines and out_lines[-1] == "":
            continue
        out_lines.append(ln)

    cleaned = "\n".join(out_lines).strip(" \n\t:-")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def _extract_cs_section(text: str, heading_patterns: List[str], stop_patterns: List[str]) -> str:
    found = _find_first_match(heading_patterns, text)
    if not found:
        return ""

    _, section_start, _ = found
    section_end = len(text)

    stop = _find_first_match(stop_patterns, text, start=section_start)
    if stop:
        section_end = stop[0]

    section = text[section_start:section_end]
    return _clean_cs_section_text(section)


def _extract_cs_background_and_summary_from_text(text: str) -> Tuple[str, str]:
    if not text.strip():
        return "", ""

    numbered = r"(?:\[\d{3,4}\]\s*)?(?:\d+\s*)?(?:[A-Z]\.\s*)?"

    background_patterns = [
        rf"(?im)^\s*{numbered}BACKGROUND\s+OF\s+THE\s+INVENTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}BACKGROUND\s+OF\s+INVENTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}BACKGROUND\s*[:\-]?\s*",
    ]
    summary_patterns = [
        rf"(?im)^\s*{numbered}SUMMARY\s+OF\s+THE\s+INVENTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}SUMMARY\s+OF\s+INVENTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}SUMMARY\s*[:\-]?\s*",
    ]
    object_patterns = [
        rf"(?im)^\s*{numbered}OBJECT(?:S)?\s+OF\s+THE\s+INVENTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}OBJECTIVE(?:S)?\s+OF\s+THE\s+INVENTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}OBJECT\s+OF\s+INVENTION\s*[:\-]?\s*",
    ]
    summary_stop_patterns = [
        rf"(?im)^\s*{numbered}BRIEF\s+DESCRIPTION(?:\s+OF\s+DRAWINGS?)?\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}DETAILED\s+DESCRIPTION(?:\s+OF\s+THE\s+INVENTION)?\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}DESCRIPTION\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}CLAIMS?\s*[:\-]?\s*",
        rf"(?im)^\s*{numbered}ABSTRACT\s*[:\-]?\s*",
    ]
    background_stop_patterns = object_patterns + summary_patterns + summary_stop_patterns

    background = _extract_cs_section(text, background_patterns, background_stop_patterns)
    summary = _extract_cs_section(text, summary_patterns, summary_stop_patterns)
    return background, summary


def extract_cs_background_and_summary(path: str) -> Tuple[str, str]:
    return _extract_cs_background_and_summary_from_text(read_pdf_text(path) or "")


def extract_cs_background_and_summary_from_docx(path: str) -> Tuple[str, str]:
    return _extract_cs_background_and_summary_from_text(read_docx_text(path) or "")


def parse_fer_pdf(path: str) -> FerParseResult:
    raw = read_pdf_text(path)
    text = _clean(raw)

    meta = _extract_meta(text)
    arts = _extract_prior_arts(text)

    detailed_obs = extract_detailed_observations_block(text)
    splits = _split_objections(detailed_obs or text)

    objections = []
    for i, (head, body) in enumerate(splits, 1):
        objections.append(
            Objection(
                number=i,
                heading=head.title(),
                body=body.strip(),
                sections=_sections_from_text(body),
                claims=_first_match(r"Claim\(s\)\s*\(([^)]+)\)", body, re.I)
                or _first_match(r"Claims?\s*[:\-]?\s*([0-9,\-\s]+)", body, re.I),
                prior_arts=arts,
            )
        )

    return FerParseResult(
        application_no=meta["application_no"],
        filing_date=meta["filing_date"],
        fer_dispatch_date=meta["fer_dispatch_date"],
        applicant=meta["applicant"],
        title=meta["title"],
        controller_name=meta["controller_name"],
        examiner_name=meta["examiner_name"],
        reply_deadline=meta["reply_deadline"],
        prior_arts=arts,
        objections=objections,
    )


def to_dict(res: FerParseResult) -> Dict:
    return asdict(res)
